"""Document parsing service ported from Junas document.rs."""
from __future__ import annotations
import io
from dataclasses import dataclass

@dataclass
class ParsedDocument:
    filename: str
    text: str
    page_count: int
    char_count: int

def parse_pdf(data: bytes, filename: str = "document.pdf") -> ParsedDocument:
    import pdfplumber
    text_parts: list[str] = []
    with pdfplumber.open(io.BytesIO(data)) as pdf:
        page_count = len(pdf.pages)
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            text_parts.append(page_text)
    text = "\n\n".join(text_parts)
    return ParsedDocument(filename=filename, text=text, page_count=page_count, char_count=len(text))

def parse_docx(data: bytes, filename: str = "document.docx") -> ParsedDocument:
    import docx
    doc = docx.Document(io.BytesIO(data))
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    return ParsedDocument(filename=filename, text=text, page_count=1, char_count=len(text))

def parse_document(data: bytes, filename: str) -> ParsedDocument:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return parse_pdf(data, filename)
    if lower.endswith(".docx"):
        return parse_docx(data, filename)
    raise ValueError(f"Unsupported file type: {filename}")
